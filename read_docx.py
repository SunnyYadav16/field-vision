
import docx
import sys

try:
    doc = docx.Document("FieldVision_Enhancement_Tasks.docx")
    print("Document Content:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"{i}: {para.text}")
    
    print("\nTables:")
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join([cell.text for cell in row.cells]))
            
except Exception as e:
    print(f"Error: {e}")
