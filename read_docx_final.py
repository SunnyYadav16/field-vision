
import docx
import sys

try:
    doc = docx.Document("FieldVision_Enhancement_Tasks.docx")
    with open('task_details.txt', 'w', encoding='utf-8') as f:
        # Just print everything but with line breaks
        for p in doc.paragraphs:
            if p.text.strip():
                f.write(p.text.strip() + "\n")
                
        f.write("\n--- TABLES ---\n")
        for table in doc.tables:
            for row in table.rows:
                # Handle possible multi-paragraph cells
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    row_cells.append(cell_text)
                
                row_text = " | ".join(row_cells)
                if row_text.strip():
                    f.write(row_text + "\n")
    print("Done")
except Exception as e:
    print(f"Error: {e}")
