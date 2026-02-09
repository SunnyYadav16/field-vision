
import docx
import sys

try:
    doc = docx.Document("FieldVision_Enhancement_Tasks.docx")
    print("Document Structure:")
    
    # Just print everything but with line breaks
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text.strip())
            
    print("\n--- TABLES ---")
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                print(row_text)
            
except Exception as e:
    print(f"Error: {e}")
